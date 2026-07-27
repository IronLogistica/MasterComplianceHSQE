"""Applica il marchio 'MasterComplianceHSQE by Maurizio Gustinicchi Consulting'
a tutti i documenti .docx/.xlsx della libreria (footer di pagina + proprieta' documento).

Da eseguire una tantum sui file gia' presenti in app/static/library/.
"""
import os
import sys

from docx import Document
from openpyxl import load_workbook

BRAND = "MasterComplianceHSQE by Maurizio Gustinicchi Consulting"
AUTHOR = "Maurizio Gustinicchi Consulting"
LIBRARY_ROOT = os.path.join(os.path.dirname(__file__), "app", "static", "library")


def brand_docx(path):
    doc = Document(path)
    for section in doc.sections:
        footer = section.footer
        footer.is_linked_to_previous = False
        # rimuove eventuali marchi gia' presenti (idempotenza), poi lo aggiunge in coda
        for p in list(footer.paragraphs):
            if p.text.strip() == BRAND:
                p._element.getparent().remove(p._element)
        new_p = footer.add_paragraph()
        new_p.text = BRAND
    try:
        doc.core_properties.comments = BRAND
        # sovrascrive sempre autore/ultimo autore: elimina ogni traccia del
        # fornitore originale (es. "Winple Italia", "MY ADVISOR SRL S.B.")
        doc.core_properties.author = AUTHOR
        doc.core_properties.last_modified_by = AUTHOR
    except Exception:
        pass
    doc.save(path)


def brand_xlsx(path):
    wb = load_workbook(path)
    for ws in wb.worksheets:
        ws.oddFooter.center.text = BRAND
        ws.oddFooter.center.size = 8
        # rimuove eventuale protezione struttura/foglio: i file devono restare modificabili
        ws.protection.sheet = False
    if wb.security is not None:
        wb.security.lockStructure = False
        wb.security.lockWindows = False
    try:
        wb.properties.company = "Maurizio Gustinicchi Consulting"
        wb.properties.description = BRAND
        wb.properties.creator = AUTHOR
        wb.properties.lastModifiedBy = AUTHOR
    except Exception:
        pass
    wb.save(path)


def main():
    ok, failed = 0, []
    for root, _dirs, files in os.walk(LIBRARY_ROOT):
        for filename in files:
            full_path = os.path.join(root, filename)
            ext = os.path.splitext(filename)[1].lower()
            try:
                if ext == ".docx":
                    brand_docx(full_path)
                    ok += 1
                elif ext == ".xlsx":
                    brand_xlsx(full_path)
                    ok += 1
            except Exception as exc:
                failed.append((full_path, str(exc)))
    print(f"Marchiati {ok} documenti.")
    if failed:
        print(f"Falliti {len(failed)}:")
        for path, err in failed:
            print(" -", path, "->", err)


if __name__ == "__main__":
    main()
